"""Word 文档处理器 (.docx)

使用 python-docx 进行内容提取，核心特性：
  1. 按段落提取（识别标题层级 Heading 1/2/3）
  2. 表格提取：读取每个表格的行和单元格，转 Markdown 表格
  3. 分片策略：按"标题+段落组+表格"分组，保留文档结构

注意：仅支持 .docx（Office 2007+）格式，不支持旧版 .doc 二进制格式。
"""

import re
from pathlib import Path
from typing import Dict, List, Optional

from langchain_core.documents import Document
from loguru import logger

from app.processors.base_processor import BaseProcessor


class WordProcessor(BaseProcessor):
    """Word (.docx) 文档处理器"""

    supported_extensions = [".docx"]

    # 标题样式匹配（用于识别 Heading 层级）
    _HEADING_PATTERN = re.compile(
        r"^Heading\s*(\d+)$", re.IGNORECASE
    )

    def read_file(self, file_path: str) -> List[Dict]:
        """读取 Word 文档，提取结构化内容

        返回结构（按内容块顺序）：
        [
            {"type": "heading", "level": 1, "content": "章节标题"},
            {"type": "paragraph", "content": "段落文本"},
            {"type": "table", "content": "| a | b |\n| --- | --- |\n| 1 | 2 |"},
            ...
        ]
        """
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise ImportError(
                "python-docx 未安装，请运行: pip install python-docx"
            ) from e

        path = Path(file_path)
        doc = DocxDocument(str(path))
        blocks: List[Dict] = []

        # 用于追踪当前段落序号
        para_index = 0

        # 遍历段落和表格（按文档中的出现顺序）
        # 注意：docx 中段落和表格是分开存储的，这里用位置推断顺序
        # 简化策略：先取所有段落，再取所有表格

        # Step 1: 提取段落（识别标题）
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 识别标题层级
            heading_level = self._extract_heading_level(para)
            if heading_level:
                blocks.append(
                    {
                        "type": "heading",
                        "level": heading_level,
                        "content": text,
                        "paragraph_index": para_index,
                    }
                )
            else:
                blocks.append(
                    {
                        "type": "paragraph",
                        "content": text,
                        "paragraph_index": para_index,
                    }
                )
            para_index += 1

        # Step 2: 提取表格（转 Markdown 表格）
        for table_idx, table in enumerate(doc.tables):
            try:
                table_data: List[List[str]] = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # 移除单元格内的换行
                        cell_text = cell_text.replace("\n", " ").replace("\r", " ")
                        row_data.append(cell_text)
                    table_data.append(row_data)

                markdown_table = self._table_to_markdown(table_data)
                if markdown_table.strip():
                    blocks.append(
                        {
                            "type": "table",
                            "content": markdown_table,
                            "table_index": table_idx,
                        }
                    )
            except Exception as e:
                logger.warning(f"Word 文档表格 {table_idx} 解析失败: {e}")

        logger.info(
            f"Word 读取完成: {file_path}, 共 {len(blocks)} 个内容块 "
            f"(段落: {sum(1 for b in blocks if b['type'] in ('heading','paragraph'))}, "
            f"表格: {sum(1 for b in blocks if b['type'] == 'table')})"
        )
        return blocks

    def split_content(self, content: List[Dict], file_path: str) -> List[Document]:
        """对 Word 内容进行分片

        策略：
          1. 按 Heading 层级做第一级分割（章节级）
          2. 每个章节内的段落累积到 chunk_size，超出则新建分片
          3. 表格作为完整独立分片（不与正文混排，避免表格被截断）
          4. 标题 metadata 记录 h1/h2/h3
        """
        documents: List[Document] = []

        # 状态：当前章节的标题信息
        current_h1: Optional[str] = None
        current_h2: Optional[str] = None
        current_h3: Optional[str] = None

        # 当前分片累积的文本
        current_segments: List[str] = []
        current_size = 0

        def flush_current():
            """将当前累积的内容输出为一个分片"""
            nonlocal current_segments, current_size
            if not current_segments:
                return

            text = "\n\n".join(current_segments).strip()
            if not text:
                current_segments = []
                current_size = 0
                return

            metadata = self._build_metadata(
                block_type="paragraph",
                extra={k: v for k, v in [
                    ("h1", current_h1),
                    ("h2", current_h2),
                    ("h3", current_h3),
                ] if v},
            )
            documents.append(Document(page_content=text, metadata=metadata))
            current_segments = []
            current_size = 0

        for block in content:
            block_type = block["type"]
            block_content = block["content"]

            # 标题：更新章节信息，作为新分片起点
            if block_type == "heading":
                level = block.get("level", 1)
                if level == 1:
                    current_h1 = block_content
                    current_h2 = None
                    current_h3 = None
                elif level == 2:
                    current_h2 = block_content
                    current_h3 = None
                elif level >= 3:
                    current_h3 = block_content

                # 标题处新建分片（标题作为新分片的第一行）
                flush_current()
                heading_line = f"{'#' * level} {block_content}"
                current_segments.append(heading_line)
                current_size = len(heading_line)
                continue

            # 表格：独立分片
            if block_type == "table":
                flush_current()
                table_metadata = self._build_metadata(
                    block_type="table",
                    extra={k: v for k, v in [
                        ("h1", current_h1),
                        ("h2", current_h2),
                        ("h3", current_h3),
                    ] if v},
                )
                # 如果表格过大，二次分割
                if len(block_content) <= self.chunk_size:
                    documents.append(
                        Document(page_content=block_content, metadata=table_metadata)
                    )
                else:
                    sub_docs = self._text_splitter.create_documents(
                        texts=[block_content], metadatas=[table_metadata]
                    )
                    documents.extend(sub_docs)
                continue

            # 普通段落：累积到当前分片
            if current_size + len(block_content) > self.chunk_size and current_segments:
                flush_current()

            current_segments.append(block_content)
            current_size += len(block_content)

        # 刷新最后一个分片
        flush_current()

        # 合并过小的分片 + 二次分割（对过长的段落分片做处理）
        return self._secondary_split(documents, min_size=200)

    # ============= 内部方法 =============

    def _extract_heading_level(self, paragraph) -> Optional[int]:
        """从段落样式中提取标题层级

        返回 1/2/3/... 表示 Heading 1/2/3/...，
        如果不是标题则返回 None。
        """
        try:
            style_name = paragraph.style.name if paragraph.style else ""
        except Exception:
            return None

        if not style_name:
            return None

        # 直接匹配 "Heading N" / "标题 N" 样式
        match = self._HEADING_PATTERN.match(style_name)
        if match:
            return int(match.group(1))

        # 中文样式名兼容（如 Word 中文版中的 "标题 1"）
        chinese_match = re.match(r"^标题\s*(\d+)$", style_name)
        if chinese_match:
            return int(chinese_match.group(1))

        # 兼容一些自定义样式名（如 "H1"、"Level 1"）
        h_match = re.match(r"^H(\d+)$", style_name, re.IGNORECASE)
        if h_match:
            return int(h_match.group(1))

        return None

    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """将二维表格数据转 Markdown 表格格式"""
        if not table_data or not table_data[0]:
            return ""

        # 过滤完全空的行
        cleaned = [
            [cell.strip() for cell in row]
            for row in table_data
        ]
        cleaned = [row for row in cleaned if any(row)]
        if not cleaned:
            return ""

        # 对齐每行列数
        col_count = max(len(row) for row in cleaned)
        for row in cleaned:
            while len(row) < col_count:
                row.append("")

        # 构建 Markdown 表格
        header = cleaned[0]
        separator = ["---"] * col_count
        rows = cleaned[1:]

        lines = []
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(separator) + " |")
        for row in rows:
            escaped = [c.replace("|", "\\|") for c in row]
            lines.append("| " + " | ".join(escaped) + " |")

        return "\n".join(lines)